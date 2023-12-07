import os
import tkinter as tk
from tkinter import messagebox
from tkinter import ttk
from PIL import Image, ImageTk
import csv
from datetime import datetime  # Import module 'datetime'
import requests
import json
from docx import Document
from tkinter import filedialog

# Hàm để xử lý khi bấm nút "Lấy dữ liệu"
# Dữ liệu lịch sử được lưu trong danh sách
lich_su_data = []
data_string = '{ "fullName": "Bành Thị Nòi", "cardIssueDate": "27/09/2021", "placeOfOrigin": "Mường Tè Lai Châu", "placeOfResident": "Giảng Võ, Ba Đình, Hà Nội", "ethnic": "Mường", "religion": "Không", "fatherName": "Bành Vân Chung", "motherName": "Phương", "spouse": "Phương Thất Phật", "identifyCharacteristics": "Nốt ruồi cách mép 2cm", "expiredDate": "02/11/2027", "documentNumber": "001202031357", "dayOfBirth": "02/11/2002", "sex": "Nam", "type": "CitizenInfo" }'
custom_font = ("Arial", 16)
# Biến đánh dấu trang hiện tại
current_page = 1
items_per_page = 10
total_pages = len(lich_su_data) // items_per_page + 1
tim_kiem_entry = None
file_path = ''
# File print
# Replace 'path/to/your/document.docx' with the actual path to your Word document
# file_path = 'don_mo_tai_khoan.docx'
# printer_name = 'don_mo_tai_khoan'


find_text = [
            '[HOVATEN]',
            '[NGAYSINH]',
            '[NGAYCAP]',
            '[NGAYHETHAN]',
            '[CCCD]',
            '[DANTOC]',
            '[TONGIAO]',
            '[HOKHAU]',
            '[QUEQUAN]',
            '[GIOITINH]'
        ]
found_row_cell = []
doc = None
def read_docx():
    global file_path
    global found_row_cell
    global doc
    file_path = filedialog.askopenfilename(title="Select a File", filetypes=[("Word files", "*.docx")])
    if file_path:
        # Process the selected file path (e.g., display it, use it in your logic)
        print("Selected File:", file_path)
        messagebox.showinfo("Thông báo", f"Đang tiến hành kiểm tra tài liệu, vui lòng chờ trong giây lát.")
        doc = Document(file_path)
    for table_index, table in enumerate(doc.tables):
        check = []
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                replace_flag = False
                for find_text_index, find_text_item in enumerate(find_text):
                    if find_text_item in cell.text:
                        # Append the indices and the index of find_text in find_text list

                        if find_text_item not in check:
                            check.append(find_text_item)
                            found_row_cell.append([table_index, row_index, cell_index, find_text_index])
                        else:
                            replace_flag = True
    messagebox.showinfo("Tạo biểu mẫu thành công", f"Đã tìm thấy {len(found_row_cell)} giá trị trong biểu mẫu")


def importdoc():
    data = get_data_from_url()
    # data = data_string
    if doc is None:
        messagebox.showerror('Lỗi', 'Bạn chưa nhập biểu mẫu')
        return
    if data is not None:
        data_json = json.loads(data)
        output_path = 'tai_khoan_'+ data_json['documentNumber'] +'.docx'

        replace_text = [
            data_json['fullName'],
            data_json['dayOfBirth'],
            data_json['cardIssueDate'],
            data_json['expiredDate'],
            data_json['documentNumber'],
            data_json['ethnic'],
            data_json['religion'],
            data_json['placeOfResident'],
            data_json['placeOfOrigin'],
            data_json['sex']
        ]

        for found_table, found_row, found_cell, find_text_index in found_row_cell:
            doc.tables[found_table].cell(found_row, found_cell).text = doc.tables[found_table].cell(found_row, found_cell).text.replace(find_text[find_text_index] ,replace_text[find_text_index])

        # for table in doc.tables:
        #     for row_index, row in table.rows:
        #         for cell_index, cell in row.cells:
        #             for find, replace in zip(find_text, replace_text):
        #                 if find_text in cell.text:
        #                     cell.text = cell.text.replace(find, replace)
        doc.save(output_path)
        os.startfile(output_path)


def get_data_from_url():
    try:
        response = requests.get('http://localhost:8689')
        response.raise_for_status()  # Nếu có lỗi HTTP, nó sẽ nâng cao một ngoại lệ
        data = response.text
        return data
    except requests.exceptions.RequestException as e:
        messagebox.showerror("Lỗi", e)
        return None

def lay_thoi_gian_hien_tai():
    now = datetime.now()
    return now.strftime("%H:%M:%S %Y-%m-%d")

# Hàm để chuyển đến trang tiếp theo
def update_page(page_curent, page):
    page_curent.config(text="Trang: " + str(page))

def trang_tiep_theo(table, page_curent):
    global current_page
    current_page += 1
    update_page(page_curent, current_page)
    keyword = tim_kiem_entry.get() if tim_kiem_entry else None
    hien_thi_lich_su(table, keyword)


# Hàm để quay lại trang trước
def trang_truoc(table, page_curent):
    global current_page
    current_page -= 1
    update_page(page_curent, current_page)
    if current_page < 1:
        current_page = 1
    keyword = tim_kiem_entry.get() if tim_kiem_entry else None
    hien_thi_lich_su(table, keyword)

# Hàm tìm kiếm
def tim_kiem(table):
    keyword = tim_kiem_entry.get().lower()  # Lấy từ khóa từ ô tìm kiếm và chuyển thành chữ thường
    hien_thi_lich_su(table, keyword)

# Hàm để hiển thị thông tin lịch sử
def hien_thi_lich_su(table, keyword=None):
    global current_page
    global items_per_page
    for item in table.get_children():
        table.delete(item)
    start = (current_page - 1) * items_per_page
    end = start + items_per_page
    for row in lich_su_data[start:end]:
        if keyword is None or any(keyword in field.lower() for field in row):
            table.insert("", "end", values=row)

# Hàm để tạo trang danh sách lịch sử

def tao_trang_lich_su():
    trang_lich_su = tk.Toplevel(root)
    trang_lich_su.title("Lịch sử")
    # Khung tìm kiếm được gán bằng biến toàn cục
    global tim_kiem_entry
    # global current_page
    background_image = Image.open("background.png")
    background_photo = ImageTk.PhotoImage(background_image)
    background_label = tk.Label(trang_lich_su, image=background_photo)
    background_label.place(x=0, y=0, relheight=1)
    # Tạo khung tìm kiếm
    khung_tim_kiem = tk.Frame(trang_lich_su)

    khung_tim_kiem.pack(fill="x")
    quay_lai_button = tk.Button(khung_tim_kiem, text="Quay lại", command=trang_lich_su.destroy, padx=10, height=2)

    quay_lai_button.pack(side="left", padx=30, pady=20)
    tim_kiem_entry = tk.Entry(khung_tim_kiem)
    tim_kiem_entry.config(font=custom_font)
    tim_kiem_entry.pack(side="left", padx=5, pady=20)
    tim_kiem_button = tk.Button(khung_tim_kiem, text="Tìm", command=lambda: tim_kiem(lich_su_table), padx=10, height=2, width=5)
    tim_kiem_button.pack(side="left", pady=20)
    page_curent = tk.Label(khung_tim_kiem, text="Trang: " + str(current_page), bg="white", fg="black", width=15)
    page_curent.pack(side="right", padx=30, pady=20)
    # Tạo nút quay lại

    # Tạo bảng dữ liệu sử dụng ttk.Treeview
    lich_su_frame = tk.Frame(trang_lich_su)
    lich_su_frame.pack(fill="both", expand=True)
    x_scrollbar = ttk.Scrollbar(lich_su_frame, orient="horizontal")
    y_scrollbar = ttk.Scrollbar(lich_su_frame, orient="vertical")
    lich_su_table = ttk.Treeview(trang_lich_su, columns=("HoTen", "GioiTinh", "SoCCCD", "NgaySinh", "HoKhau", "DacDiem", "QueQuan", "NgayCapCCCD", "ThoiGian", "NguoiCanGap", "PhongBan", "MucDich"), xscrollcommand=x_scrollbar.set, yscrollcommand=y_scrollbar.set)
    x_scrollbar.config(command=lich_su_table.xview)
    y_scrollbar.config(command=lich_su_table.yview)
    # Đặt tiêu đề cột
    lich_su_table.heading("#0", text="", anchor="center")
    lich_su_table.heading("HoTen", text="Họ tên")
    lich_su_table.heading("GioiTinh", text="Giới tính")
    lich_su_table.heading("SoCCCD", text="Số CCCD")
    lich_su_table.heading("NgaySinh", text="Ngày sinh")
    lich_su_table.heading("HoKhau", text="Hộ khẩu thường trú")
    lich_su_table.heading("DacDiem", text="Đặc điểm nhận dạng")
    lich_su_table.heading("QueQuan", text="Quê quán")
    lich_su_table.heading("NgayCapCCCD", text="Ngày cấp CCCD")
    lich_su_table.heading("ThoiGian", text="Thời gian")
    lich_su_table.heading("NguoiCanGap", text="Note1")
    lich_su_table.heading("PhongBan", text="Note2")
    lich_su_table.heading("MucDich", text="Note3")
    lich_su_table.pack(fill="both", expand=True)
    # Đặt chiều rộng cột
    lich_su_table.column("#0", width=6)
    lich_su_table.column("HoTen", width=150)
    lich_su_table.column("GioiTinh", width=60)
    lich_su_table.column("SoCCCD", width=90)
    lich_su_table.column("NgaySinh", width=80)
    lich_su_table.column("HoKhau", width=110)
    lich_su_table.column("DacDiem", width=120)
    lich_su_table.column("QueQuan", width=100)
    lich_su_table.column("NgayCapCCCD", width=100)
    lich_su_table.column("ThoiGian", width=120)
    lich_su_table.column("NguoiCanGap", width=40)
    lich_su_table.column("PhongBan", width=40)
    lich_su_table.column("MucDich", width=40)

    # Set the row height for data rows

    # Hiển thị bảng dữ liệu và các nút

    lich_su_table.pack()
    # Tạo nút trang trước và trang tiếp theo
    trang_truoc_button = tk.Button(trang_lich_su, text="Trang trước", command=lambda: trang_truoc(lich_su_table, page_curent))
    trang_tiep_theo_button = tk.Button(trang_lich_su, text="Trang tiếp",  command=lambda: trang_tiep_theo(lich_su_table, page_curent))

    trang_truoc_button.pack(side="left", padx=30, pady=20)
    trang_tiep_theo_button.pack(side="right", padx=30, pady=20)

    # Đọc và hiển thị dữ liệu từ tệp CSV
    doc_lich_su()
    hien_thi_lich_su(lich_su_table)
    style = ttk.Style()
    style.configure('Treeview', rowheight=40)  # Adjust the height as needed for the header row


# Hàm để đọc thông tin lịch sử từ tệp CSV
def doc_lich_su():
    global lich_su_data
    lich_su_data = []  # Khởi tạo biến lich_su_data
    try:
        with open("lich_su.csv", mode="r", encoding='utf-8') as file:
            reader = csv.reader(file)
            for row in reader:
                lich_su_data.append(row)
    except FileNotFoundError:
        lich_su_data = []


def lay_du_lieu():
    data = get_data_from_url()
    if data is not None:
        data_json = json.loads(data)
        # Lấy dữ liệu từ các trường nhập liệu
        with open('lich_su.csv', 'a', newline='', encoding='utf-8') as file:
            writer = csv.writer(file)
            writer.writerow(
                [data_json["fullName"], data_json["sex"], data_json["documentNumber"], data_json["dayOfBirth"], data_json["placeOfResident"], data_json["identifyCharacteristics"], data_json["placeOfOrigin"], data_json["cardIssueDate"], lay_thoi_gian_hien_tai(), '', '', ''])

            messagebox.showinfo("Thông báo", "Đã lưu thành công!")
    else:
        messagebox.showerror("Lỗi", "Không tìm thấy dữ liệu")
    #     data_json = json.loads(data_string)
    #     ho_ten_entry.insert(0, data_json["fullName"])
    #     gioi_tinh_entry.insert(0, data_json["sex"])
    #     so_cccd_entry.insert(0, data_json["documentNumber"])
    #     ngay_sinh_entry.insert(0, data_json["dayOfBirth"])
    #     ho_khau_entry.insert(0, data_json["placeOfResident"])
    #     dac_diem_entry.insert(0, data_json["identifyCharacteristics"])
    #     que_quan_entry.insert(0, data_json["placeOfOrigin"])
    #     ngay_cap_cccd_entry.insert(0, data_json["cardIssueDate"])
    #     with open('lich_su.csv', 'a', newline='', encoding='utf-8') as file:
    #         writer = csv.writer(file)
    #         writer.writerow(
    #             [data_json["fullName"], data_json["sex"], data_json["documentNumber"], data_json["dayOfBirth"], data_json["placeOfResident"], data_json["identifyCharacteristics"], data_json["placeOfOrigin"], data_json["cardIssueDate"], lay_thoi_gian_hien_tai(), '', '', ''])
    #
    #         messagebox.showinfo("Thông báo", "Đã lưu thành công!")



# Tạo cửa sổ chính
root = tk.Tk()
root.title("Quản lý lịch sử")

# root.geometry("800x600")


background_image = Image.open("background.png")
# Chuyển đổi đối tượng Image thành đối tượng PhotoImage
background_photo = ImageTk.PhotoImage(background_image)

# Tạo một Label với hình nền là "abc.jpg"
background_label = tk.Label(root, image=background_photo)
background_label.place(x=0, y=0, relheight=1)



lich_su_button = tk.Button(root, text="Lịch sử", command=tao_trang_lich_su, padx=10, borderwidth=0, height=2)
lich_su_button.grid(row=0, column=0, padx=10, pady=10, sticky='w')


# Tạo nút "Lấy dữ liệu"
lay_du_lieu_button = tk.Button(root, text="Lấy dữ liệu", command=lay_du_lieu, padx=10, borderwidth=0,  height=2)
lay_du_lieu_button.grid(row=0, column=1, padx=(10, 10), pady=10, sticky='w')


read_doc = tk.Button(root, text="Chọn biểu mẫu", command=read_docx, padx=10, borderwidth=0, height=2)
read_doc.grid(row=1, column=0, padx=(10, 10), pady=10, sticky='w')
# Tạo nút "Import"
import_doc = tk.Button(root, text="Xuất biểu mẫu", command=importdoc, padx=10, borderwidth=0,  height=2)
import_doc.grid(row=1, column=1, padx=(10, 10), pady=10, sticky='w')



labels = ["Họ tên:", "Ngày sinh:", "Số CCCD:", "Đặc điểm ND:", "Hộ khẩu TT:", "Ngày cấp CCCD:", "Người cần gặp:", "Phòng ban:", "Mục đích:"]
entries = []





# Bắt đầu chạy ứng dụng
root.mainloop()